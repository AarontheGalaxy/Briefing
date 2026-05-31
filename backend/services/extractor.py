from typing import Tuple
import io


MAX_FILE_SIZE_MB = 50
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024


async def extract_text(file_content: bytes, filename: str) -> Tuple[str, int]:
    if len(file_content) > MAX_FILE_SIZE_BYTES:
        raise ValueError(f"File exceeds the {MAX_FILE_SIZE_MB}MB limit.")

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        text = _extract_pdf(file_content)
    elif ext == "docx":
        text = _extract_docx(file_content)
    elif ext in ("txt", "md"):
        try:
            text = file_content.decode("utf-8").strip()
        except UnicodeDecodeError:
            # Fall back to latin-1 which never fails on arbitrary bytes
            text = file_content.decode("latin-1").strip()
    else:
        raise ValueError("Only PDF, DOCX, and TXT files are supported.")

    word_count = len(text.split())
    return text, word_count


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages: list[str] = []

    for page in reader.pages:
        page_text = page.extract_text() or ""
        stripped = page_text.strip()
        if stripped:
            pages.append(stripped)

    return "\n\n".join(pages)


def _extract_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)
